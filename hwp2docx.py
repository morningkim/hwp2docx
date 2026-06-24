# -*- coding: utf-8 -*-
"""
hwp2docx.py — HWP 5.0(바이너리)를 직접 파싱하여 DOCX로 변환

오피스 프로그램(한컴/MS/LibreOffice) 없이 순수 Python으로 동작.
의존성: olefile, python-docx  (zlib, struct 는 표준 라이브러리)

처리 범위:
  - 여러 Section(BodyText/Section0..N)
  - 문단 본문 텍스트(제어문자 폭 처리: char=1 / inline·extended=8 워드)
  - 표: 행·열, 셀 병합(colSpan/rowSpan), 셀 내부 문단, 중첩 표는 텍스트로 평탄화
한계:
  - 글꼴/글자 크기/색상/정밀 레이아웃/이미지/도형은 재현하지 않음(텍스트·표 구조 위주)

사용법:
  python hwp2docx.py                 # 현재 폴더의 모든 *.hwp 변환
  python hwp2docx.py a.hwp           # a.hwp -> a.docx
  python hwp2docx.py a.hwp out.docx  # 출력 경로 지정
"""

import sys
import os
import glob
import zlib
import struct

import olefile
from docx import Document

# ---- HWP 레코드 태그 ----
HWPTAG_BEGIN = 0x10
PARA_HEADER  = HWPTAG_BEGIN + 50   # 66
PARA_TEXT    = HWPTAG_BEGIN + 51   # 67
CTRL_HEADER  = HWPTAG_BEGIN + 55   # 71
LIST_HEADER  = HWPTAG_BEGIN + 56   # 72
TABLE        = HWPTAG_BEGIN + 61   # 77

# 제어문자 폭: char 컨트롤(1워드)
CHAR_CONTROLS = {0, 10, 13, 24, 25, 26, 27, 28, 29, 30, 31}
# 나머지 제어문자(1~31 중 위 집합 제외)는 inline/extended = 8워드


def u16(buf, off):
    if off + 2 > len(buf):
        return 0
    return struct.unpack_from('<H', buf, off)[0]


def parse_records(data):
    """섹션 스트림 -> [(tag, level, payload), ...]"""
    recs = []
    i, n = 0, len(data)
    while i + 4 <= n:
        header = struct.unpack_from('<I', data, i)[0]
        i += 4
        tag = header & 0x3FF
        level = (header >> 10) & 0x3FF
        size = (header >> 20) & 0xFFF
        if size == 0xFFF:
            size = struct.unpack_from('<I', data, i)[0]
            i += 4
        payload = data[i:i + size]
        i += size
        recs.append((tag, level, payload))
    return recs


class Node:
    __slots__ = ('tag', 'level', 'payload', 'children')

    def __init__(self, tag, level, payload):
        self.tag = tag
        self.level = level
        self.payload = payload
        self.children = []


def build_tree(records):
    """level 기반으로 레코드를 트리로 구성"""
    roots = []
    stack = []
    for tag, level, payload in records:
        node = Node(tag, level, payload)
        while stack and stack[-1].level >= level:
            stack.pop()
        if stack:
            stack[-1].children.append(node)
        else:
            roots.append(node)
        stack.append(node)
    return roots


def decode_text(payload):
    """PARA_TEXT payload(UTF-16LE + 제어문자) -> 문자열"""
    n = len(payload) // 2
    if n == 0:
        return ''
    arr = struct.unpack_from('<%dH' % n, payload, 0)
    out = []
    i = 0
    while i < n:
        c = arr[i]
        if c < 32:
            if c == 9:            # tab (inline, 8워드)
                out.append('\t')
                i += 8
            elif c in (10, 13):   # 줄/문단 나눔 (char, 1워드)
                out.append('\n')
                i += 1
            elif c in CHAR_CONTROLS:
                i += 1            # 기타 char 컨트롤: 표시 텍스트 없음
            else:
                i += 8            # inline/extended 컨트롤: 8워드 건너뜀
        else:
            out.append(chr(c))
            i += 1
    return ''.join(out)


def build_table(ctrl_node):
    """CTRL_HEADER('tbl ') 노드 -> 표 블록

    셀 문단(PARA_HEADER)이 LIST_HEADER의 자식으로 들어가는 경우와,
    형제로 이어지는 경우(파일/버전에 따라 레벨 인코딩이 다름)를 모두 처리한다.
    """
    tbl = {'type': 'table', 'nRows': 0, 'nCols': 0, 'cells': []}
    current = None
    for ch in ctrl_node.children:
        if ch.tag == TABLE:
            p = ch.payload
            tbl['nRows'] = u16(p, 4)
            tbl['nCols'] = u16(p, 6)
        elif ch.tag == LIST_HEADER:
            p = ch.payload
            # 셀 속성 오프셋(실측): col=8, row=10, colSpan=12, rowSpan=14
            current = {
                'col':     u16(p, 8),
                'row':     u16(p, 10),
                'colSpan': u16(p, 12),
                'rowSpan': u16(p, 14),
                'nodes':   list(ch.children),  # 자식으로 들어간 셀 내용
            }
            tbl['cells'].append(current)
        elif current is not None:
            # LIST_HEADER 뒤에 형제로 이어진 셀 내용
            current['nodes'].append(ch)

    for cell in tbl['cells']:
        cell['blocks'] = build_blocks(cell.pop('nodes'))
    return tbl


def build_blocks(nodes):
    """노드 리스트 -> 문서 블록(문단/표) 리스트 (문서 순서 유지)"""
    blocks = []
    for node in nodes:
        if node.tag == PARA_HEADER:
            text = ''
            tables = []
            for ch in node.children:
                if ch.tag == PARA_TEXT:
                    text += decode_text(ch.payload)
                elif ch.tag == CTRL_HEADER:
                    ctrl_id = ch.payload[0:4][::-1]
                    if ctrl_id == b'tbl ':
                        tables.append(build_table(ch))
            t = text.replace('\x00', '').strip()
            if t:
                blocks.append({'type': 'para', 'text': t})
            blocks.extend(tables)
    return blocks


def blocks_to_text(blocks):
    """블록(중첩 표 포함)을 텍스트로 평탄화 (셀 내부용 폴백)"""
    out = []
    for b in blocks:
        if b['type'] == 'para':
            out.append(b['text'])
        elif b['type'] == 'table':
            for cell in b['cells']:
                txt = blocks_to_text(cell['blocks'])
                if txt:
                    out.append(txt)
    return '\n'.join(x for x in out if x)


def add_table_to_doc(doc, tb):
    nRows, nCols = tb['nRows'], tb['nCols']
    cells = tb['cells']
    # 행/열 정보가 없거나 좌표가 격자를 벗어나면 텍스트로 폴백
    if nRows <= 0 or nCols <= 0:
        for cell in cells:
            txt = blocks_to_text(cell['blocks'])
            if txt:
                doc.add_paragraph(txt)
        return

    table = doc.add_table(rows=nRows, cols=nCols)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass

    # 1) 셀 텍스트를 좌표에 먼저 배치
    for cell in cells:
        r, c = cell['row'], cell['col']
        if 0 <= r < nRows and 0 <= c < nCols:
            table.cell(r, c).text = blocks_to_text(cell['blocks'])

    # 2) 병합: 덮을 영역이 비어 있을 때만(내용 손실 방지) 병합
    for cell in cells:
        r, c = cell['row'], cell['col']
        rs = max(1, cell['rowSpan'])
        cs = max(1, cell['colSpan'])
        if (rs == 1 and cs == 1) or r >= nRows or c >= nCols:
            continue
        r2 = min(nRows - 1, r + rs - 1)
        c2 = min(nCols - 1, c + cs - 1)
        conflict = False
        for rr in range(r, r2 + 1):
            for cc in range(c, c2 + 1):
                if (rr, cc) != (r, c) and table.cell(rr, cc).text.strip():
                    conflict = True
                    break
            if conflict:
                break
        if conflict:
            continue
        try:
            table.cell(r, c).merge(table.cell(r2, c2))
        except Exception:
            pass


def hwp_to_blocks(path):
    ole = olefile.OleFileIO(path)
    try:
        fh = ole.openstream('FileHeader').read()
        if fh[0:17] != b'HWP Document File':
            raise ValueError('HWP 5.0 형식이 아닙니다: %s' % path)
        flags = struct.unpack_from('<I', fh, 36)[0]
        compressed = bool(flags & 0x01)
        encrypted = bool(flags & 0x02)
        if encrypted:
            raise ValueError('암호로 보호된 문서는 지원하지 않습니다: %s' % path)

        # 섹션 수집/정렬
        sections = []
        for entry in ole.listdir():
            if len(entry) == 2 and entry[0] == 'BodyText' and entry[1].startswith('Section'):
                try:
                    num = int(entry[1][len('Section'):])
                except ValueError:
                    num = 0
                sections.append((num, entry))
        sections.sort(key=lambda x: x[0])

        all_blocks = []
        for _, entry in sections:
            data = ole.openstream(entry).read()
            if compressed:
                data = zlib.decompress(data, -15)
            recs = parse_records(data)
            tree = build_tree(recs)
            all_blocks.extend(build_blocks(tree))
        return all_blocks
    finally:
        ole.close()


def convert(src, dst=None):
    if dst is None:
        dst = os.path.splitext(src)[0] + '.docx'
    blocks = hwp_to_blocks(src)
    doc = Document()
    for b in blocks:
        if b['type'] == 'para':
            doc.add_paragraph(b['text'])
        elif b['type'] == 'table':
            add_table_to_doc(doc, b)
    doc.save(dst)
    n_tbl = sum(1 for b in blocks if b['type'] == 'table')
    n_par = sum(1 for b in blocks if b['type'] == 'para')
    print('  -> %s  (문단 %d, 표 %d)' % (dst, n_par, n_tbl))
    return dst


def main(argv):
    if len(argv) >= 2:
        src = argv[1]
        dst = argv[2] if len(argv) >= 3 else None
        print('변환:', src)
        convert(src, dst)
    else:
        files = sorted(glob.glob('*.hwp'))
        if not files:
            print('현재 폴더에 *.hwp 파일이 없습니다.')
            return
        for f in files:
            print('변환:', f)
            try:
                convert(f)
            except Exception as e:
                print('  [실패]', e)


if __name__ == '__main__':
    main(sys.argv)
